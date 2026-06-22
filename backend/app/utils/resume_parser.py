import io
import re


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text content from a PDF file using PyPDF2 with enhanced extraction."""
    from PyPDF2 import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    
    full_text = "\n".join(text_parts).strip()
    
    # Clean up common PDF extraction artifacts
    if full_text:
        # Fix broken words from column extraction
        full_text = re.sub(r'(\w)-\n(\w)', r'\1\2', full_text)
        # Fix excessive whitespace
        full_text = re.sub(r' {3,}', '  ', full_text)
        # Fix lines that are just whitespace
        full_text = re.sub(r'\n\s*\n\s*\n', '\n\n', full_text)
    
    return full_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from a DOCX file."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    
    # Extract from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())
    
    # Also extract from tables (many resumes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n".join(text_parts).strip()


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract text from uploaded file based on extension. Multiple fallback strategies."""
    lower = filename.lower()
    extracted_text = ""
    
    # Strategy 1: Use the appropriate parser based on file extension
    try:
        if lower.endswith(".pdf"):
            extracted_text = extract_text_from_pdf(file_bytes)
        elif lower.endswith(".docx"):
            extracted_text = extract_text_from_docx(file_bytes)
        elif lower.endswith(".txt") or lower.endswith(".md"):
            extracted_text = file_bytes.decode("utf-8", errors="ignore").strip()
    except Exception as e:
        print(f"[ResumeParser] Primary parser failed for {filename}: {e}")
    
    # If we got meaningful text (> 30 chars), return it
    if extracted_text and len(extracted_text.strip()) > 30:
        print(f"[ResumeParser] Successfully extracted {len(extracted_text)} chars from {filename}")
        return extracted_text
    
    # Strategy 2: Try reading raw bytes as UTF-8 text (works for some PDFs and text files)
    try:
        raw_text = file_bytes.decode("utf-8", errors="ignore").strip()
        # Filter out binary garbage - check if most characters are printable
        printable_ratio = sum(1 for c in raw_text[:500] if c.isprintable() or c in '\n\r\t') / max(len(raw_text[:500]), 1)
        if printable_ratio > 0.7 and len(raw_text) > 30:
            print(f"[ResumeParser] Fallback UTF-8 extraction got {len(raw_text)} chars from {filename}")
            return raw_text
    except Exception:
        pass
    
    # Strategy 3: Try latin-1 encoding
    try:
        raw_text = file_bytes.decode("latin-1", errors="ignore").strip()
        printable_ratio = sum(1 for c in raw_text[:500] if c.isprintable() or c in '\n\r\t') / max(len(raw_text[:500]), 1)
        if printable_ratio > 0.7 and len(raw_text) > 30:
            print(f"[ResumeParser] Fallback Latin-1 extraction got {len(raw_text)} chars from {filename}")
            return raw_text
    except Exception:
        pass
    
    # If everything fails, return whatever we got (even if short) 
    # DO NOT return hardcoded fake resume data
    if extracted_text:
        print(f"[ResumeParser] WARNING: Only extracted {len(extracted_text)} chars from {filename}")
        return extracted_text
    
    print(f"[ResumeParser] ERROR: Could not extract any text from {filename}")
    return f"[Resume file: {filename} - text extraction failed. File size: {len(file_bytes)} bytes]"
